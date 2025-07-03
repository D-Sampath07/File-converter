from docx import Document

def txt_to_docx(txt_path, output_path):
    if not txt_path.lower().endswith(".txt"):
        raise ValueError("Input file must be a TXT file.")

    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        doc = Document()
        for line in lines:
            line = line.strip()
            if line:
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")  # Preserve blank lines

        doc.save(output_path)

    except Exception as e:
        raise RuntimeError(f"Failed to convert TXT to DOCX: {e}")

    return output_path
